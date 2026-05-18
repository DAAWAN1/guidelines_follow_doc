import fitz  # PyMuPDF
import pdfplumber
import io
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# ---------------------------------------------------------------------------
# Alignment constant map (handle both int and enum values safely)
# ---------------------------------------------------------------------------
_ALIGN_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    WD_ALIGN_PARAGRAPH.DISTRIBUTE: "justify",
    None: None,  # inherit from style
}

# Common sans-serif font families (lowercase for comparison)
SANS_SERIF_FONTS = {
    "arial", "helvetica", "calibri", "verdana", "tahoma",
    "trebuchet ms", "trebuchet", "franklin gothic", "gill sans",
    "open sans", "lato", "roboto", "noto sans", "segoe ui",
    "myriad pro", "futura", "century gothic", "optima",
    "lucida sans", "lucida grande", "ubuntu", "fira sans",
    "source sans pro", "nunito", "raleway", "poppins", "inter",
    "avenir", "proxima nova", "gotham", "brandon grotesque",
    "barlow", "dm sans", "work sans", "outfit",
}


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _resolve_alignment(para) -> str:
    """Walk the style inheritance chain to resolve paragraph alignment."""
    # Direct paragraph format wins
    if para.alignment is not None:
        return _ALIGN_MAP.get(para.alignment, "left")

    # Walk up the style chain
    style = para.style
    while style:
        fmt = style.paragraph_format
        if fmt and fmt.alignment is not None:
            return _ALIGN_MAP.get(fmt.alignment, "left")
        style = style.base_style

    return "left"  # Word default


def _resolve_font_size(run, para) -> float | None:
    """Return font size in pt, walking the style inheritance chain."""
    if run.font.size:
        return round(run.font.size.pt, 1)

    style = para.style
    while style:
        if style.font and style.font.size:
            return round(style.font.size.pt, 1)
        style = style.base_style

    return None


def _resolve_font_name(run, para) -> str | None:
    """Return font name, walking the style inheritance chain."""
    if run.font.name:
        return run.font.name

    style = para.style
    while style:
        if style.font and style.font.name:
            return style.font.name
        style = style.base_style

    return None


def _resolve_bold(run, para) -> bool:
    """Resolve bold, accounting for explicit False overrides."""
    if run.bold is not None:
        return bool(run.bold)
    style = para.style
    while style:
        if style.font and style.font.bold is not None:
            return bool(style.font.bold)
        style = style.base_style
    return False


def _resolve_italic(run, para) -> bool:
    """Resolve italic, accounting for explicit False overrides."""
    if run.italic is not None:
        return bool(run.italic)
    style = para.style
    while style:
        if style.font and style.font.italic is not None:
            return bool(style.font.italic)
        style = style.base_style
    return False


def _resolve_all_caps(run, para) -> bool:
    """Resolve all_caps, accounting for style inheritance."""
    if run.font.all_caps is not None:
        return bool(run.font.all_caps)
    style = para.style
    while style:
        if style.font and style.font.all_caps is not None:
            return bool(style.font.all_caps)
        style = style.base_style
    return False


def _is_list_paragraph(para) -> bool:
    """Return True if this paragraph is formatted as a list."""
    style_name = (para.style.name or "").lower()
    if "list" in style_name:
        return True
    # Check for numPr (numbering properties) in the XML
    pPr = para._p.find(qn("w:pPr"))
    if pPr is not None and pPr.find(qn("w:numPr")) is not None:
        return True
    return False


def _get_bullet_char(para) -> str | None:
    """
    Try to extract the actual bullet character for list paragraphs.
    Returns the character or None if it cannot be determined.
    """
    style_name = (para.style.name or "").lower()
    if "list bullet" in style_name:
        return "•"   # Word default bullet
    if "list number" in style_name:
        return "1."  # Numbered list
    # For direct-formatted lists, look at the first non-empty run
    for run in para.runs:
        t = run.text.strip()
        if t and len(t) == 1 and not t.isalnum():
            return t
    return None


def _get_hyperlinks(para) -> list[dict]:
    """
    Extract hyperlinks from a paragraph's XML.
    Returns a list of dicts with keys: display_text, url.
    """
    hyperlinks = []
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
          "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships"}

    for hl in para._p.findall(".//w:hyperlink", ns):
        # Get the display text from all child runs
        display_text = "".join(
            r.text or ""
            for r in hl.findall(".//w:r/w:t", ns)
        )
        # Relationship ID → URL is in the part's relationships
        r_id = hl.get(qn("r:id"))
        url = None
        if r_id:
            try:
                part = para.part
                rel = part.rels.get(r_id)
                if rel:
                    url = rel.target_ref
            except Exception:
                pass
        if display_text or url:
            hyperlinks.append({"display_text": display_text.strip(), "url": url or ""})

    return hyperlinks


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def extract_text(uploaded_file):
    """
    Extract text from a PDF or DOCX file.
    Returns text: str.
    """
    pdf_bytes = uploaded_file.read()
    extension = (
        uploaded_file.name.rsplit(".", 1)[-1].lower()
        if "." in uploaded_file.name
        else "pdf"
    )

    if extension == "pdf":
        text = ""
        try:
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            for page in doc:
                text += page.get_text()
            doc.close()
            if text.strip():
                return text
        except Exception:
            pass

        try:
            with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                text = "\n".join(page.extract_text() or "" for page in pdf.pages)
            return text
        except Exception:
            return ""

    elif extension == "docx":
        try:
            doc = Document(io.BytesIO(pdf_bytes))
            return "\n".join(para.text for para in doc.paragraphs)
        except Exception:
            return ""
    else:
        return ""

def _run_from_xml(run_elem, para, default_font, default_size):
    """
    Extract run properties from a <w:r> element (including hyperlink runs).
    Returns a dict with the same keys as the existing run dicts.
    """
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    # ---- Text ----
    text_parts = [t.text for t in run_elem.findall(".//w:t", ns) if t.text]
    text = "".join(text_parts)

    # ---- Formatting (direct run properties) ----
    rPr = run_elem.find("w:rPr", ns)
    bold = False
    italic = False
    underline = False
    all_caps = False
    font_name = default_font
    font_size = default_size
    color_rgb = None

    if rPr is not None:
        bold = rPr.find("w:b", ns) is not None
        italic = rPr.find("w:i", ns) is not None
        underline = rPr.find("w:u", ns) is not None
        all_caps = rPr.find("w:caps", ns) is not None

        # Font name
        font_elem = rPr.find("w:rFonts", ns)
        if font_elem is not None:
            font_name = (font_elem.get(qn("w:ascii")) or
                         font_elem.get(qn("w:hAnsi")) or
                         default_font)

        # Font size (half-points → points)
        sz_elem = rPr.find("w:sz", ns)
        if sz_elem is not None:
            sz_val = sz_elem.get(qn("w:val"))
            if sz_val:
                font_size = round(int(sz_val) / 2, 1)

        # Colour
        color_elem = rPr.find("w:color", ns)
        if color_elem is not None:
            color_rgb = color_elem.get(qn("w:val"))

    # ---- Sans-serif detection ----
    is_sans_serif = None
    if font_name:
        font_lower = font_name.lower()
        is_sans_serif = (font_lower.split()[0] in SANS_SERIF_FONTS or
                         font_lower in SANS_SERIF_FONTS)

    return {
        "text": text,
        "bold": bold,
        "italic": italic,
        "underline": underline,
        "all_caps": all_caps,
        "font_name": font_name,
        "font_size_pt": font_size,
        "color_rgb": color_rgb,
        "is_sans_serif": is_sans_serif,
    }

def extract_rich_docx_data(file_source) -> dict:
    """
    Extract rich formatting data from a DOCX file for compliance checking.

    Parameters
    ----------
    file_source : bytes | file-like object
        Raw bytes of the .docx file, or any file-like object readable
        by python-docx's Document().

    Returns
    -------
    dict with keys:
        paragraphs : list[dict]   – one dict per paragraph (see below)
        page_count : int          – estimated page count (word-count heuristic)
        has_toc    : bool         – True if a Table of Contents field was found
        default_font : str | None – document default font name
        default_size : float | None – document default font size in pt

    Paragraph dict keys:
        text        : str           – plain text
        style       : str           – style name (e.g. "Heading 1", "Normal")
        is_heading  : bool
        is_title    : bool          – True for "Title" style
        is_caption  : bool          – True for "Caption" style
        alignment   : str           – "left" | "center" | "right" | "justify"
        is_list     : bool
        bullet_char : str | None
        hyperlinks  : list[dict]    – [{display_text, url}, …]
        runs        : list[dict]    – one dict per run (see below)

    Run dict keys:
        text        : str
        bold        : bool
        italic      : bool
        underline   : bool
        all_caps    : bool
        font_name   : str | None
        font_size_pt: float | None
        color_rgb   : str | None    – hex string e.g. "FF0000", or None
        is_sans_serif: bool | None  – None if font_name unknown
    """
    if isinstance(file_source, (bytes, bytearray)):
        file_source = io.BytesIO(file_source)

    doc = Document(file_source)

    # ---- Document-level defaults ----------------------------------------
    default_font = None
    default_size = None
    try:
        normal_style = doc.styles.get_by_name("Normal")
        if normal_style:
            if normal_style.font.name:
                default_font = normal_style.font.name
            if normal_style.font.size:
                default_size = round(normal_style.font.size.pt, 1)
    except Exception:
        pass

    # ---- Table of Contents detection ------------------------------------
    has_toc = False
    body_xml = doc.element.body.xml
    if "TOC" in body_xml or "w:fldChar" in body_xml:
        # Check for TOC field code
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        for instr in doc.element.body.findall(".//w:instrText", ns):
            if instr.text and "TOC" in instr.text.upper():
                has_toc = True
                break

    # ---- Paragraph extraction -------------------------------------------
    paragraphs_data = []

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else "Normal"
        style_lower = style_name.lower()
        is_heading = style_lower.startswith("heading")
        is_title = style_lower == "title"
        is_caption = "caption" in style_lower

        alignment = _resolve_alignment(para)
        is_list = _is_list_paragraph(para)
        bullet_char = _get_bullet_char(para) if is_list else None
        hyperlinks = _get_hyperlinks(para)

        runs_data = []
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        for run_elem in para._p.findall(".//w:r", ns):
            run_dict = _run_from_xml(run_elem, para, default_font, default_size)
            runs_data.append(run_dict)

        paragraphs_data.append({
            "text": para.text,
            "style": style_name,
            "is_heading": is_heading,
            "is_title": is_title,
            "is_caption": is_caption,
            "alignment": alignment,
            "is_list": is_list,
            "bullet_char": bullet_char,
            "hyperlinks": hyperlinks,
            "runs": runs_data,
        })

    # ---- Estimate page count (300 words ≈ 1 page) -----------------------
    total_words = sum(len(p["text"].split()) for p in paragraphs_data)
    page_count = max(1, round(total_words / 300))

    # Build full plain text from all paragraphs (for text‑only checks)
    full_text = "\n".join(p["text"] for p in paragraphs_data)

    return {
        "paragraphs": paragraphs_data,
        "page_count": page_count,
        "has_toc": has_toc,
        "default_font": default_font,
        "default_size": default_size,
        "full_text": full_text,   # <-- new line
    }